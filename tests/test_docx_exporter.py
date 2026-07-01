"""
Test DOCX Exporter

Purpose:
Unit tests for deterministic DOCX export rendering.
"""

import copy

from docx import Document

from engine.docx_exporter import export_resume_to_docx


def export_payload() -> dict:
    """Return a valid resume export payload fixture."""
    return {
        "status": "success",
        "errors": [],
        "document_title": "Administrative Assistant",
        "sections": [
            {
                "id": "professional_summary",
                "title": "Professional Summary",
                "type": "paragraph",
                "content": (
                    "Organized administrative professional with experience in "
                    "documentation and reporting."
                ),
            },
            {
                "id": "skills",
                "title": "Skills",
                "type": "grouped_list",
                "content": {
                    "Technical": ["Microsoft Excel", "Documentation"],
                    "Soft Skills": ["Communication"],
                    "Tools": [],
                    "Domain": ["Administrative Support"],
                },
            },
            {
                "id": "experience",
                "title": "Experience",
                "type": "bullet_list",
                "content": [
                    "Maintained Excel trackers and prepared administrative reports."
                ],
            },
        ],
        "metadata": {
            "adapter": "resume_export_adapter",
            "version": "v1",
            "source": "final_resume",
            "source_metadata": {},
        },
    }


def paragraph_text(output_path) -> list[str]:
    """Return text from all DOCX paragraphs."""
    return [paragraph.text for paragraph in Document(output_path).paragraphs]


def test_valid_export_payload_creates_docx_file(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"

    result = export_resume_to_docx(export_payload(), str(output_path))

    assert result["status"] == "success"
    assert output_path.exists()


def test_return_shape_contains_required_fields(tmp_path) -> None:
    result = export_resume_to_docx(export_payload(), str(tmp_path / "resume.docx"))

    assert set(result) == {"status", "errors", "output_path", "metadata"}
    assert set(result["metadata"]) == {"exporter", "version", "source"}


def test_docx_file_exists_after_successful_export(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"

    export_resume_to_docx(export_payload(), str(output_path))

    assert output_path.is_file()


def test_created_docx_can_be_reopened(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"

    export_resume_to_docx(export_payload(), str(output_path))

    assert Document(output_path).paragraphs


def test_created_docx_contains_expected_text(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"

    export_resume_to_docx(export_payload(), str(output_path))
    text = "\n".join(paragraph_text(output_path))

    assert "Administrative Assistant" in text
    assert "Professional Summary" in text
    assert "Organized administrative professional" in text
    assert "Skills" in text
    assert "Technical" in text
    assert "Microsoft Excel" in text
    assert "Experience" in text
    assert "Maintained Excel trackers" in text


def test_non_dict_export_payload_returns_failed(tmp_path) -> None:
    result = export_resume_to_docx([], str(tmp_path / "resume.docx"))

    assert result["status"] == "failed"
    assert result["errors"] == ["export_payload must be a dictionary."]


def test_failed_export_payload_status_returns_failed(tmp_path) -> None:
    payload = export_payload()
    payload["status"] = "failed"

    result = export_resume_to_docx(payload, str(tmp_path / "resume.docx"))

    assert result["status"] == "failed"
    assert "export_payload must have status 'success'." in result["errors"]


def test_empty_output_path_returns_failed() -> None:
    result = export_resume_to_docx(export_payload(), "")

    assert result["status"] == "failed"
    assert "output_path must be a non-empty string." in result["errors"]


def test_missing_or_non_list_sections_returns_failed(tmp_path) -> None:
    payload = export_payload()
    payload["sections"] = {}

    result = export_resume_to_docx(payload, str(tmp_path / "resume.docx"))

    assert result["status"] == "failed"
    assert "export_payload sections must be a list." in result["errors"]


def test_paragraph_section_renders_without_crashing(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"
    payload = export_payload()
    payload["sections"] = [payload["sections"][0]]

    result = export_resume_to_docx(payload, str(output_path))

    assert result["status"] == "success"
    assert "Professional Summary" in paragraph_text(output_path)


def test_grouped_skills_section_renders_only_non_empty_groups(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"

    export_resume_to_docx(export_payload(), str(output_path))
    text = "\n".join(paragraph_text(output_path))

    assert "Technical" in text
    assert "Soft Skills" in text
    assert "Domain" in text
    assert "Tools" not in text


def test_bullet_list_section_renders_bullet_items(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"

    export_resume_to_docx(export_payload(), str(output_path))

    assert any(
        "Maintained Excel trackers" in text for text in paragraph_text(output_path)
    )


def test_empty_content_does_not_create_placeholder_text(tmp_path) -> None:
    output_path = tmp_path / "resume.docx"
    payload = export_payload()
    payload["document_title"] = ""
    payload["sections"] = [
        {
            "id": "professional_summary",
            "title": "Professional Summary",
            "type": "paragraph",
            "content": "",
        }
    ]

    export_resume_to_docx(payload, str(output_path))
    text = "\n".join(paragraph_text(output_path))

    assert "Professional Summary" in text
    assert "No content" not in text
    assert "Placeholder" not in text


def test_unsupported_section_type_returns_failed_with_clear_error(tmp_path) -> None:
    payload = export_payload()
    payload["sections"].append(
        {"id": "bad", "title": "Bad", "type": "unknown", "content": []}
    )

    result = export_resume_to_docx(payload, str(tmp_path / "resume.docx"))

    assert result["status"] == "failed"
    assert "Unsupported section type: unknown" in result["errors"]


def test_parent_output_directory_is_created_if_missing(tmp_path) -> None:
    output_path = tmp_path / "exports" / "resume.docx"

    result = export_resume_to_docx(export_payload(), str(output_path))

    assert result["status"] == "success"
    assert output_path.exists()


def test_output_path_with_no_parent_directory_is_handled_safely(tmp_path, monkeypatch) -> None:
    monkeypatch.chdir(tmp_path)

    result = export_resume_to_docx(export_payload(), "resume.docx")

    assert result["status"] == "success"
    assert (tmp_path / "resume.docx").exists()


def test_exporter_does_not_mutate_input_payload(tmp_path) -> None:
    payload = export_payload()
    original = copy.deepcopy(payload)

    export_resume_to_docx(payload, str(tmp_path / "resume.docx"))

    assert payload == original
