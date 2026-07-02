"""
Test Resume Text Extractor

Purpose:
Unit tests for deterministic DOCX resume text extraction.
"""

from pathlib import Path

from docx import Document

from pypdf import PdfWriter

from engine.resume_text_extractor import extract_text_from_docx, extract_text_from_pdf


def write_docx(path: Path, paragraphs: list[str] | None = None) -> None:
    """Create a small DOCX fixture."""
    document = Document()
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)
    document.save(path)


def _pdf_string(value: str) -> str:
    """Escape a value for a simple PDF text string."""
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def write_text_pdf(path: Path, pages: list[list[str]]) -> None:
    """Create a small selectable-text PDF fixture."""
    objects: list[str] = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    page_refs: list[str] = []
    next_object_id = 4

    for page_lines in pages:
        page_id = next_object_id
        content_id = next_object_id + 1
        next_object_id += 2
        page_refs.append(f"{page_id} 0 R")
        lines = ["BT", "/F1 12 Tf", "72 720 Td"]
        for index, line in enumerate(page_lines):
            if index:
                lines.append("0 -18 Td")
            lines.append(f"({_pdf_string(line)}) Tj")
        lines.append("ET")
        stream = "\n".join(lines)
        objects.append(
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_id} 0 R >>"
        )
        objects.append(
            f"<< /Length {len(stream.encode('ascii'))} >>\n"
            f"stream\n{stream}\nendstream"
        )

    objects[1] = f"<< /Type /Pages /Kids [{' '.join(page_refs)}] /Count {len(page_refs)} >>"

    offsets: list[int] = []
    body = b"%PDF-1.4\n"
    for object_id, content in enumerate(objects, start=1):
        offsets.append(len(body))
        body += f"{object_id} 0 obj\n{content}\nendobj\n".encode("ascii")

    xref_offset = len(body)
    xref_entries = ["0000000000 65535 f "]
    xref_entries.extend(f"{offset:010d} 00000 n " for offset in offsets)
    xref = "\n".join(xref_entries)
    trailer = (
        f"xref\n0 {len(objects) + 1}\n{xref}\n"
        f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    )
    path.write_bytes(body + trailer.encode("ascii"))


def test_extracts_docx_paragraph_text(tmp_path) -> None:
    docx_path = tmp_path / "resume.docx"
    write_docx(
        docx_path,
        [
            "Sample Candidate",
            "Administrative Assistant",
            "Maintained Excel trackers and prepared reports.",
        ],
    )

    result = extract_text_from_docx(str(docx_path))

    assert result["status"] == "success"
    assert result["text"] == (
        "Sample Candidate\n"
        "Administrative Assistant\n"
        "Maintained Excel trackers and prepared reports."
    )
    assert result["errors"] == []
    assert result["warnings"] == []


def test_extracts_docx_table_text(tmp_path) -> None:
    docx_path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Sample Candidate")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Microsoft Excel"
    table.cell(1, 0).text = "Experience"
    table.cell(1, 1).text = "Coordinated office documentation"
    document.save(docx_path)

    result = extract_text_from_docx(str(docx_path))

    assert result["status"] == "success"
    assert result["text"].splitlines() == [
        "Sample Candidate",
        "Skills",
        "Microsoft Excel",
        "Experience",
        "Coordinated office documentation",
    ]


def test_return_shape_contains_status_text_errors_warnings_metadata(tmp_path) -> None:
    docx_path = tmp_path / "resume.docx"
    write_docx(docx_path, ["Sample Candidate"])

    result = extract_text_from_docx(str(docx_path))

    assert set(result) == {"status", "text", "errors", "warnings", "metadata"}
    assert set(result["metadata"]) == {"extractor", "version", "source"}


def test_omits_empty_paragraphs_and_cells(tmp_path) -> None:
    docx_path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Sample Candidate")
    document.add_paragraph("")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = ""
    table.cell(0, 1).text = "Reporting"
    document.save(docx_path)

    result = extract_text_from_docx(str(docx_path))

    assert result["text"].splitlines() == ["Sample Candidate", "Reporting"]


def test_normalizes_internal_whitespace(tmp_path) -> None:
    docx_path = tmp_path / "resume.docx"
    write_docx(docx_path, ["  Managed\t schedules   and\u00a0reports  "])

    result = extract_text_from_docx(str(docx_path))

    assert result["text"] == "Managed schedules and reports"


def test_non_string_file_path_returns_failed() -> None:
    result = extract_text_from_docx(Path("resume.docx"))

    assert result["status"] == "failed"
    assert result["text"] == ""
    assert result["errors"] == ["file_path must be a string."]
    assert result["warnings"] == []


def test_empty_file_path_returns_failed() -> None:
    result = extract_text_from_docx("")

    assert result["status"] == "failed"
    assert result["errors"] == ["file_path must be a non-empty string."]


def test_non_docx_extension_returns_failed(tmp_path) -> None:
    text_path = tmp_path / "resume.txt"
    text_path.write_text("Sample Candidate", encoding="utf-8")

    result = extract_text_from_docx(str(text_path))

    assert result["status"] == "failed"
    assert result["errors"] == ["file_path must point to a .docx file."]


def test_missing_docx_file_returns_failed(tmp_path) -> None:
    result = extract_text_from_docx(str(tmp_path / "missing.docx"))

    assert result["status"] == "failed"
    assert result["errors"] == ["file_path file does not exist."]


def test_directory_path_returns_failed(tmp_path) -> None:
    directory_path = tmp_path / "resume.docx"
    directory_path.mkdir()

    result = extract_text_from_docx(str(directory_path))

    assert result["status"] == "failed"
    assert result["errors"] == ["file_path must point to a file."]


def test_empty_docx_returns_failed(tmp_path) -> None:
    docx_path = tmp_path / "empty.docx"
    write_docx(docx_path)

    result = extract_text_from_docx(str(docx_path))

    assert result["status"] == "failed"
    assert result["errors"] == ["DOCX contains no extractable text."]


def test_corrupt_docx_returns_failed_with_clear_error(tmp_path) -> None:
    docx_path = tmp_path / "corrupt.docx"
    docx_path.write_text("not a real docx", encoding="utf-8")

    result = extract_text_from_docx(str(docx_path))

    assert result["status"] == "failed"
    assert result["errors"][0].startswith("DOCX text extraction failed:")


def test_pdf_return_shape_matches_extractor_contract(tmp_path) -> None:
    pdf_path = tmp_path / "resume.pdf"
    write_text_pdf(pdf_path, [["Sample Candidate"]])

    result = extract_text_from_pdf(str(pdf_path))

    assert set(result) == {"status", "text", "errors", "warnings", "metadata"}
    assert result["status"] == "success"
    assert result["errors"] == []
    assert result["warnings"] == []
    assert result["metadata"]["page_count"] == 1


def test_extracts_single_page_selectable_text_pdf(tmp_path) -> None:
    pdf_path = tmp_path / "resume.pdf"
    write_text_pdf(
        pdf_path,
        [["Sample Candidate", "Administrative Assistant", "Microsoft Excel"]],
    )

    result = extract_text_from_pdf(str(pdf_path))

    assert result["status"] == "success"
    assert "Sample Candidate" in result["text"]
    assert "Administrative Assistant" in result["text"]
    assert "Microsoft Excel" in result["text"]


def test_extracts_multi_page_selectable_text_pdf(tmp_path) -> None:
    pdf_path = tmp_path / "resume.pdf"
    write_text_pdf(
        pdf_path,
        [
            ["Sample Candidate", "Administrative Assistant"],
            ["Experience", "Prepared reports"],
        ],
    )

    result = extract_text_from_pdf(str(pdf_path))

    assert result["status"] == "success"
    assert result["metadata"]["page_count"] == 2
    assert "Sample Candidate" in result["text"]
    assert "Prepared reports" in result["text"]
    assert "\n\n" in result["text"]


def test_blank_pdf_returns_failed_result(tmp_path) -> None:
    pdf_path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with pdf_path.open("wb") as pdf_file:
        writer.write(pdf_file)

    result = extract_text_from_pdf(str(pdf_path))

    assert result["status"] == "failed"
    assert result["text"] == ""
    assert result["errors"] == ["PDF contains no extractable selectable text."]
    assert result["metadata"]["page_count"] == 1


def test_missing_pdf_path_returns_failed_result(tmp_path) -> None:
    result = extract_text_from_pdf(str(tmp_path / "missing.pdf"))

    assert result["status"] == "failed"
    assert result["errors"] == ["file_path file does not exist."]


def test_invalid_pdf_returns_failed_result(tmp_path) -> None:
    pdf_path = tmp_path / "invalid.pdf"
    pdf_path.write_text("not a real pdf", encoding="utf-8")

    result = extract_text_from_pdf(str(pdf_path))

    assert result["status"] == "failed"
    assert result["errors"][0].startswith("PDF text extraction failed:")
