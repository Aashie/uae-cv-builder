"""
DOCX Exporter

Purpose:
Render a resume export payload into a .docx file.
"""

from pathlib import Path

from docx import Document


EXPORTER_METADATA = {
    "exporter": "docx_exporter",
    "version": "v1",
    "source": "resume_export_payload",
}


def _result(status: str, errors: list[str], output_path: str) -> dict:
    """Return the DOCX exporter result shape."""
    return {
        "status": status,
        "errors": errors,
        "output_path": output_path,
        "metadata": EXPORTER_METADATA.copy(),
    }


def _validate_export_request(export_payload: dict, output_path: str) -> list[str]:
    """Validate exporter inputs without repairing them."""
    if not isinstance(export_payload, dict):
        return ["export_payload must be a dictionary."]

    errors: list[str] = []
    if export_payload.get("status") != "success":
        errors.append("export_payload must have status 'success'.")
    if not isinstance(output_path, str) or not output_path:
        errors.append("output_path must be a non-empty string.")
    if not isinstance(export_payload.get("sections"), list):
        errors.append("export_payload sections must be a list.")

    return errors


def _add_grouped_list(document: Document, content: dict) -> None:
    """Render grouped list content into a document."""
    if not isinstance(content, dict):
        return

    for group_label, items in content.items():
        if not isinstance(items, list) or not items:
            continue
        group_paragraph = document.add_paragraph()
        group_paragraph.add_run(str(group_label)).bold = True
        for item in items:
            if item:
                document.add_paragraph(str(item), style="List Bullet")


def _add_bullet_list(document: Document, content: list) -> None:
    """Render bullet list content into a document."""
    if not isinstance(content, list):
        return

    for item in content:
        if item:
            document.add_paragraph(str(item), style="List Bullet")


def _render_section(document: Document, section: dict, errors: list[str]) -> None:
    """Render a single export payload section."""
    title = section.get("title", "") if isinstance(section, dict) else ""
    section_type = section.get("type") if isinstance(section, dict) else None
    content = section.get("content", "") if isinstance(section, dict) else ""

    if title:
        document.add_heading(title, level=1)

    if section_type == "paragraph":
        if content:
            document.add_paragraph(str(content))
    elif section_type == "grouped_list":
        _add_grouped_list(document, content)
    elif section_type == "bullet_list":
        _add_bullet_list(document, content)
    else:
        errors.append(f"Unsupported section type: {section_type}")


def export_resume_to_docx(export_payload: dict, output_path: str) -> dict:
    """Export a resume payload to a DOCX file."""
    validation_errors = _validate_export_request(export_payload, output_path)
    if validation_errors:
        return _result("failed", validation_errors, output_path)

    errors: list[str] = []
    document = Document()

    document_title = export_payload.get("document_title", "")
    if document_title:
        document.add_heading(str(document_title), level=0)

    for section in export_payload["sections"]:
        _render_section(document, section, errors)

    if errors:
        return _result("failed", errors, output_path)

    try:
        output = Path(output_path)
        if output.parent != Path("."):
            output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
    except Exception as error:
        return _result("failed", [f"DOCX export failed: {error}"], output_path)

    return _result("success", [], output_path)
